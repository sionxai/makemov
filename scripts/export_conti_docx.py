#!/usr/bin/env python3
"""
적벽대전 줄콘티 → DOCX 표 문서 변환기 v2
컬럼: 컷 | 화면묘사 | 대사 | 콘티 이미지(빈칸 — 9:16)
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 경로 설정 ──
CONTI_PATH = os.path.join(os.path.dirname(__file__), '..', 'src', 'data', 'redcliff-conti.js')
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), '..', 'output', '적벽대전_줄콘티.docx')

# ── JS 파싱 ──

def extract_field(text, field_name, start_pos):
    patterns = [
        rf"{field_name}:\s*'((?:[^'\\]|\\.)*)'",
        rf"{field_name}:\s*\"((?:[^\"\\]|\\.)*)\"",
        rf'{field_name}:\s*`((?:[^`\\]|\\.)*)`',
    ]
    for pat in patterns:
        m = re.search(pat, text[start_pos:], re.DOTALL)
        if m:
            val = m.group(1).replace('\\n', '\n').replace("\\'", "'")
            return val
    return ''

def extract_number(text, field_name, start_pos):
    m = re.search(rf"{field_name}:\s*([\d.]+)", text[start_pos:])
    return float(m.group(1)) if m else 0

def parse_scenes(content):
    title_m = re.search(r"title:\s*'([^']*)'", content)
    title = title_m.group(1) if title_m else '적벽대전 줄콘티'
    dur_m = re.search(r"totalDuration:\s*'([^']*)'", content)
    total_duration = dur_m.group(1) if dur_m else ''

    scenes = []
    scene_pattern = re.compile(r"scene_id:\s*'(S\d+)'")
    scene_positions = [(m.group(1), m.start()) for m in scene_pattern.finditer(content)]

    for i, (scene_id, spos) in enumerate(scene_positions):
        end_pos = scene_positions[i+1][1] if i+1 < len(scene_positions) else len(content)
        scene_text = content[spos:end_pos]

        heading = extract_field(content, 'heading', spos)

        cuts = []
        cut_pattern = re.compile(r"cut_id:\s*'(S\d+-C\d+)'")
        cut_positions = [(m.group(1), m.start()) for m in cut_pattern.finditer(scene_text)]

        for j, (cut_id, cpos) in enumerate(cut_positions):
            cut = {
                'cut_id': cut_id,
                'visual': extract_field(scene_text, 'visual', cpos),
                'dialogue': extract_field(scene_text, 'dialogue', cpos),
            }
            cuts.append(cut)

        scenes.append({
            'scene_id': scene_id,
            'heading': heading,
            'cuts': cuts,
        })

    return title, total_duration, scenes

# ── DOCX 유틸 ──

def set_cell_bg(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, font_name='맑은 고딕'):
    """셀에 텍스트 추가 (기존 내용 대체)"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(13)

    lines = text.split('\n') if text else ['']
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = RGBColor(*color)
        if i < len(lines) - 1:
            run.add_break()

def set_row_height(row, height_cm):
    """행 높이 고정"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height_cm * 567)}" w:hRule="atLeast"/>')
    trPr.append(trHeight)

def add_placeholder_text(cell, text, size=8):
    """이미지 자리 표시 텍스트"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 수직 가운데 정렬
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)

    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.italic = True

# ── 문서 생성 ──

def create_docx(title, total_duration, scenes):
    doc = Document()

    # A4 가로
    section = doc.sections[0]
    section.orientation = 1
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # ── 타이틀 ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)
    run = title_p.add_run(title)
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(6)
    run = sub_p.add_run(f'총 {total_duration} | 줄콘티 (Line Conti)')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # ── 컬럼 정의 ──
    headers = ['컷', '화면 묘사', '대사', '콘티 이미지']
    col_widths = [Cm(2.0), Cm(10.0), Cm(8.5), Cm(5.5)]
    IMAGE_ROW_HEIGHT = 7.5  # 9:16 비율 (너비 5.5cm 기준 → 높이 약 9.8cm, 여유 두고 7.5cm)

    HEADER_BG = '2C3E50'
    SCENE_BG = 'D4A574'
    ROW_EVEN = 'F7F3EE'
    ROW_ODD = 'FFFFFF'

    for scene in scenes:
        # ── 씬 헤더 ──
        scene_p = doc.add_paragraph()
        scene_p.paragraph_format.space_before = Pt(8)
        scene_p.paragraph_format.space_after = Pt(3)
        run = scene_p.add_run(f'  {scene["scene_id"]}  |  {scene["heading"]}')
        run.font.size = Pt(12)
        run.bold = True
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        run.font.color.rgb = RGBColor(255, 255, 255)
        # 배경색
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{SCENE_BG}"/>')
        scene_p.paragraph_format.element.get_or_add_pPr().append(shading)

        # ── 테이블 ──
        table = doc.add_table(rows=1 + len(scene['cuts']), cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # 컬럼 너비
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = w

        # 헤더 행
        for ci, h in enumerate(headers):
            cell = table.rows[0].cells[ci]
            set_cell_bg(cell, HEADER_BG)
            add_cell_text(cell, h, bold=True, size=9, color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)

        # 데이터 행
        for ri, cut in enumerate(scene['cuts']):
            row = table.rows[ri + 1]
            bg = ROW_EVEN if ri % 2 == 0 else ROW_ODD

            # 행 높이 (이미지 들어갈 공간)
            set_row_height(row, IMAGE_ROW_HEIGHT)

            # 컷 ID
            cell0 = row.cells[0]
            set_cell_bg(cell0, bg)
            add_cell_text(cell0, cut['cut_id'], bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            # 수직 가운데
            vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
            cell0._tc.get_or_add_tcPr().append(vAlign)

            # 화면 묘사
            cell1 = row.cells[1]
            set_cell_bg(cell1, bg)
            add_cell_text(cell1, cut['visual'], size=9)

            # 대사
            cell2 = row.cells[2]
            set_cell_bg(cell2, bg)
            if cut['dialogue']:
                add_cell_text(cell2, cut['dialogue'], size=9, color=(139, 69, 19))
            else:
                add_cell_text(cell2, '', size=9)

            # 콘티 이미지 (빈칸 — 플레이스홀더)
            cell3 = row.cells[3]
            set_cell_bg(cell3, 'F0F0F0')
            add_placeholder_text(cell3, f'📷\n{cut["cut_id"]}\n(9:16 이미지 삽입)', size=8)

        # 테이블 후 여백
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(2)
        spacer.paragraph_format.space_after = Pt(2)

    # ── 푸터 ──
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_p.add_run('Generated by makemov | 줄콘티 v2')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 저장
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f'✅ 문서 생성 완료: {OUTPUT_PATH}')
    print(f'   씬 {len(scenes)}개, 총 컷 {sum(len(s["cuts"]) for s in scenes)}개')

# ── 실행 ──
if __name__ == '__main__':
    content = open(CONTI_PATH, 'r', encoding='utf-8').read()
    title, total_duration, scenes = parse_scenes(content)

    print(f'📄 파싱 완료: {title} ({total_duration})')
    for s in scenes:
        print(f'   {s["scene_id"]}: {s["heading"]} — {len(s["cuts"])}컷')

    create_docx(title, total_duration, scenes)
